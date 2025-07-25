"""
Document Processor for Single Document Chat

Handles document loading, parsing, and chunking for individual documents.
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
from datetime import datetime

try:
    import PyPDF2
    import docx
    import pandas as pd
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")


class DocumentProcessor:
    """
    Processes documents for single document chat functionality.
    
    Supports multiple document formats and provides text chunking
    for efficient retrieval and processing.
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks in characters
            chunk_overlap: Overlap between chunks in characters
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        self.logger = logging.getLogger(__name__)
        
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load and parse a document from file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata and content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
            
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                content = self._load_pdf(file_path)
            elif file_extension == '.docx':
                content = self._load_docx(file_path)
            elif file_extension == '.txt':
                content = self._load_txt(file_path)
            elif file_extension in ['.csv', '.xlsx']:
                content = self._load_spreadsheet(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
            # Generate document hash for caching
            doc_hash = self._generate_document_hash(file_path)
            
            return {
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'file_extension': file_extension,
                'content': content,
                'doc_hash': doc_hash,
                'loaded_at': datetime.now().isoformat(),
                'chunk_count': 0
            }
            
        except Exception as e:
            self.logger.error(f"Error loading document {file_path}: {e}")
            raise
            
    def _load_pdf(self, file_path: Path) -> str:
        """Load content from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {e}")
            raise
            
    def _load_docx(self, file_path: Path) -> str:
        """Load content from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            self.logger.error(f"Error reading DOCX {file_path}: {e}")
            raise
            
    def _load_txt(self, file_path: Path) -> str:
        """Load content from text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            self.logger.error(f"Error reading text file {file_path}: {e}")
            raise
            
    def _load_spreadsheet(self, file_path: Path) -> str:
        """Load content from spreadsheet file."""
        try:
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Convert dataframe to text representation
            text = df.to_string(index=False)
            return text
        except Exception as e:
            self.logger.error(f"Error reading spreadsheet {file_path}: {e}")
            raise
            
    def _generate_document_hash(self, file_path: Path) -> str:
        """Generate a hash for the document content."""
        try:
            with open(file_path, 'rb') as file:
                content = file.read()
                return hashlib.md5(content).hexdigest()
        except Exception:
            # Fallback to file path hash if content reading fails
            return hashlib.md5(str(file_path).encode()).hexdigest()
            
    def chunk_document(self, document: Dict[str, Any]) -> List[Document]:
        """
        Split document content into chunks for processing.
        
        Args:
            document: Document dictionary from load_document()
            
        Returns:
            List of Document chunks with metadata
        """
        try:
            content = document['content']
            if not content.strip():
                raise ValueError("Document content is empty")
                
            # Create langchain Document objects
            docs = [Document(
                page_content=content,
                metadata={
                    'source': document['file_path'],
                    'file_name': document['file_name'],
                    'doc_hash': document['doc_hash']
                }
            )]
            
            # Split into chunks
            chunks = self.text_splitter.split_documents(docs)
            
            # Update document with chunk count
            document['chunk_count'] = len(chunks)
            
            self.logger.info(f"Created {len(chunks)} chunks from document {document['file_name']}")
            return chunks
            
        except Exception as e:
            self.logger.error(f"Error chunking document: {e}")
            raise
            
    def get_document_summary(self, document: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a summary of document metadata and statistics.
        
        Args:
            document: Document dictionary from load_document()
            
        Returns:
            Dictionary with document summary information
        """
        content = document['content']
        word_count = len(content.split())
        char_count = len(content)
        
        return {
            'file_name': document['file_name'],
            'file_size_mb': round(document['file_size'] / (1024 * 1024), 2),
            'word_count': word_count,
            'character_count': char_count,
            'chunk_count': document['chunk_count'],
            'file_type': document['file_extension'],
 